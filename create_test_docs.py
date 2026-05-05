"""Генератор тестовых DOCX-файлов с богатыми метаданными."""
from docx import Document
from datetime import datetime
import os

OUT = "test_docs/clean"
os.makedirs(OUT, exist_ok=True)


# ── 1. Криминалистический акт (нормальные метаданные) ─────
doc1 = Document()
doc1.core_properties.author           = "Иванов Иван Иванович"
doc1.core_properties.last_modified_by = "Петров Алексей"
doc1.core_properties.created          = datetime(2024, 3, 15, 9, 0)
doc1.core_properties.modified         = datetime(2024, 11, 20, 17, 33)
doc1.core_properties.revision         = 7
doc1.core_properties.title            = "Акт экспертного исследования №128"
doc1.core_properties.subject          = "Компьютерно-техническая экспертиза"

doc1.add_heading("Акт экспертного исследования №128", 0)
doc1.add_paragraph(
    "Настоящий акт составлен по результатам компьютерно-технической экспертизы "
    "цифровых носителей информации."
)
doc1.add_paragraph("Эксперт: Иванов И.И., стаж 12 лет, квалификация: судебный эксперт 1-й категории.")
doc1.add_paragraph(
    "Объекты исследования: жёсткий диск WD Blue 2TB, "
    "мобильный телефон Samsung Galaxy S21 Ultra."
)
doc1.add_heading("Выводы", 1)
doc1.add_paragraph(
    "1. На исследованном накопителе обнаружены признаки удаления файлов "
    "с помощью специализированного ПО (CCleaner v5.88)."
)
doc1.add_paragraph(
    "2. Дата последнего доступа к удалённым файлам — 14 марта 2024 года."
)
doc1.add_paragraph("Подпись эксперта: _________________")
path1 = f"{OUT}/forensic_act_128.docx"
doc1.save(path1)
print(f"✅ {path1}")


# ── 2. Договор (много правок, история редактирования) ─────
doc2 = Document()
doc2.core_properties.author           = "ООО ЮрПомощь"
doc2.core_properties.last_modified_by = "Сидорова Мария"
doc2.core_properties.created          = datetime(2023, 8, 1, 10, 0)
doc2.core_properties.modified         = datetime(2025, 1, 10, 14, 22)
doc2.core_properties.revision         = 23
doc2.core_properties.title            = "Договор оказания юридических услуг"
doc2.core_properties.subject          = "Гражданско-правовой договор"

doc2.add_heading("Договор №ЮП-2023/088", 0)
doc2.add_paragraph("г. Москва                                          01 августа 2023 г.")
doc2.add_paragraph(
    "ООО «ЮрПомощь», именуемое в дальнейшем «Исполнитель», "
    "и Гражданин Кузнецов Д.А., именуемый «Заказчик», "
    "заключили настоящий договор."
)
doc2.add_heading("1. Предмет договора", 1)
doc2.add_paragraph(
    "Исполнитель обязуется оказать юридические услуги в объёме, "
    "указанном в Приложении №1."
)
doc2.add_heading("2. Стоимость и порядок расчётов", 1)
doc2.add_paragraph("Стоимость услуг составляет 85 000 (восемьдесят пять тысяч) рублей.")
path2 = f"{OUT}/contract_legal_v23.docx"
doc2.save(path2)
print(f"✅ {path2}")


# ── 3. Документ с аномальными датами (форензик-ловушка) ───
doc3 = Document()
doc3.core_properties.author           = "root"
doc3.core_properties.last_modified_by = "Administrator"
doc3.core_properties.created          = datetime(1980, 1, 1, 0, 0)    # аномалия: до эпохи NTFS
doc3.core_properties.modified         = datetime(2077, 12, 31, 23, 59)  # аномалия: будущее
doc3.core_properties.revision         = 1
doc3.core_properties.title            = "System Configuration Manual"

doc3.add_heading("System Configuration Manual", 0)
doc3.add_paragraph("Author: root  |  Machine: WIN-SERVER-01  |  Build: 2024.01")
doc3.add_paragraph("This document describes critical system configuration parameters.")
doc3.add_paragraph("WARNING: Modification without authorization is prohibited.")
path3 = f"{OUT}/anomaly_dates_suspicious.docx"
doc3.save(path3)
print(f"✅ {path3}")


# ── 4. Отчёт с подозрительным автором (другая организация) ─
doc4 = Document()
doc4.core_properties.author           = "John Smith"
doc4.core_properties.last_modified_by = "Неизвестный пользователь"
doc4.core_properties.created          = datetime(2025, 2, 14, 3, 47)   # ночное время
doc4.core_properties.modified         = datetime(2025, 2, 14, 3, 52)
doc4.core_properties.revision         = 1
doc4.core_properties.title            = "Конфиденциальный финансовый отчёт Q4"
doc4.core_properties.subject          = "Финансы"

doc4.add_heading("Конфиденциальный финансовый отчёт Q4 2024", 0)
doc4.add_paragraph("Данный документ содержит сведения ограниченного распространения.")
doc4.add_paragraph("Итого оборот: 124 500 000 руб.")
doc4.add_paragraph("Чистая прибыль: 18 700 000 руб.")
doc4.add_paragraph("Документ создан: 14 февраля 2025, 03:47")
path4 = f"{OUT}/suspicious_author_report.docx"
doc4.save(path4)
print(f"✅ {path4}")


# ── 5. Чистый простой документ (baseline) ─────────────────
doc5 = Document()
doc5.core_properties.author           = "Студент Курсовой"
doc5.core_properties.last_modified_by = "Студент Курсовой"
doc5.core_properties.created          = datetime(2025, 5, 1, 12, 0)
doc5.core_properties.modified         = datetime(2025, 5, 1, 12, 30)
doc5.core_properties.revision         = 1
doc5.core_properties.title            = "Тестовый документ — эталон"

doc5.add_heading("Тестовый документ — эталон", 0)
doc5.add_paragraph("Этот документ является эталонным: без аномалий, без макросов.")
doc5.add_paragraph("Используется для проверки базовой работы анализатора метаданных.")
path5 = f"{OUT}/baseline_clean.docx"
doc5.save(path5)
print(f"✅ {path5}")

print("\nВсе тестовые DOCX созданы.")
